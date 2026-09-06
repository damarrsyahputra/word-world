from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from app.document.editor import (
    add_page_numbers,
    clear_page_numbers_for_sections,
    detect_dominant_font,
    to_bytes,
)


def _contains_page_field(container) -> bool:
    for instruction in container._element.iter(qn("w:instrText")):
        if "PAGE" in (instruction.text or "").upper():
            return True
    return False


def _pg_num_type(document, section_index: int):
    sect_pr = document.sections[section_index]._sectPr
    return sect_pr.find(qn("w:pgNumType"))


def test_add_page_numbers_places_field_in_footer():
    document = Document()
    add_page_numbers(document, formats=["decimal"], positions=["bottom"], alignments=["center"])
    assert _contains_page_field(document.sections[0].footer)
    assert not _contains_page_field(document.sections[0].header)


def test_add_page_numbers_places_field_in_header():
    document = Document()
    add_page_numbers(document, formats=["decimal"], positions=["top"], alignments=["right"])
    assert _contains_page_field(document.sections[0].header)
    assert not _contains_page_field(document.sections[0].footer)


def test_add_page_numbers_sets_format_roman():
    document = Document()
    add_page_numbers(document, formats=["lowerRoman"], positions=["bottom"], alignments=["center"])
    pg_num_type = _pg_num_type(document, 0)
    assert pg_num_type is not None
    assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"


def test_add_page_numbers_sets_start_number():
    document = Document()
    add_page_numbers(document, formats=["decimal"], positions=["bottom"], alignments=["center"], start_numbers=[7])
    pg_num_type = _pg_num_type(document, 0)
    assert pg_num_type is not None
    assert pg_num_type.get(qn("w:start")) == "7"


def test_add_page_numbers_none_removes_field():
    document = Document()
    add_page_numbers(document, formats=["decimal"], positions=["bottom"], alignments=["center"])
    assert _contains_page_field(document.sections[0].footer)
    add_page_numbers(document, formats=["none"], positions=["bottom"], alignments=["center"])
    assert not _contains_page_field(document.sections[0].footer)


def test_add_page_numbers_handles_multiple_sections():
    document = Document()
    document.add_section(WD_SECTION.NEW_PAGE)
    add_page_numbers(
        document,
        formats=["upperRoman", "decimal"],
        positions=["bottom", "top"],
        alignments=["center", "left"],
    )
    assert len(document.sections) == 2
    assert _contains_page_field(document.sections[0].footer)
    assert _contains_page_field(document.sections[1].header)


def test_add_page_numbers_rejects_invalid_alignment():
    document = Document()
    with pytest.raises(ValueError):
        add_page_numbers(document, alignments=["middle"])


def test_add_page_numbers_requires_one_setting_per_section():
    document = Document()
    document.add_section(WD_SECTION.NEW_PAGE)
    with pytest.raises(ValueError):
        add_page_numbers(document, formats=["decimal"])


def test_clear_page_numbers_for_sections_removes_fields():
    document = Document()
    add_page_numbers(document, formats=["decimal"], positions=["bottom"], alignments=["center"])
    assert _contains_page_field(document.sections[0].footer)
    clear_page_numbers_for_sections(document, {0})
    assert not _contains_page_field(document.sections[0].footer)


def test_to_bytes_returns_valid_docx():
    document = Document()
    document.add_paragraph("Halo dunia")
    raw = to_bytes(document)
    assert raw[:2] == b"PK"
    reloaded = Document(BytesIO(raw))
    assert reloaded.paragraphs[0].text == "Halo dunia"


def test_detect_dominant_font_zero_length_document_safe():
    document = Document()
    name, size = detect_dominant_font(document)
    assert isinstance(name, str)
    assert isinstance(size, float)