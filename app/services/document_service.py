from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
import re

from app.agent.chapters import chapter_number_from_text, parse_chapter_token
from app.agent.schemas import DocumentCommand
from app.document.editor import add_page_numbers, clear_page_numbers_for_sections, to_bytes, detect_dominant_font
from app.document.reader import list_paragraphs, load_document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass(frozen=True)
class PageNumberConflict:
    section_index: int
    ranges: tuple[str, ...]


def analyze_command(document, command: DocumentCommand) -> tuple[list[PageNumberConflict], dict[str, int]]:
    """Resolve ranges and report incompatible settings sharing one section."""
    if command.action != "configure_page_number_ranges":
        return [], {}

    paragraphs = list_paragraphs(document, include_empty=True)
    resolved: dict[str, int] = {}
    section_ranges: dict[int, list[str]] = {}
    for page_range in command.ranges:
        start_section = _find_anchor_section(paragraphs, page_range.start_anchor)
        end_section = (
            len(document.sections) - 1
            if page_range.end_anchor == "__DOCUMENT_END__"
            else _find_anchor_section(paragraphs, page_range.end_anchor)
        )
        if start_section > end_section:
            raise ValueError(
                f"Anchor range is reversed: {page_range.start_anchor!r} -> {page_range.end_anchor!r}"
            )
        resolved[page_range.start_anchor] = start_section
        resolved[page_range.end_anchor] = end_section
        label = f"{page_range.start_anchor} -> {page_range.end_anchor}"
        for section_index in range(start_section, end_section + 1):
            section_ranges.setdefault(section_index, []).append(label)

    conflicts = [
        PageNumberConflict(section_index, tuple(labels))
        for section_index, labels in section_ranges.items()
        if len(labels) > 1
    ]
    return conflicts, resolved


def _is_main_body(paragraph) -> bool:
    """True if the paragraph is in the main document body, False if in a table or textbox."""
    parent = paragraph._p.getparent()
    return parent is not None and parent.tag.endswith("body")

def _find_next_isolation_index(document, paragraphs, end_idx: int) -> int | None:
    """Find the index of the next major chapter or boundary after end_idx."""
    end_para = paragraphs[end_idx]
    
    current_num = chapter_number_from_text(_normalize_anchor(end_para.text))
    is_heading = "heading" in end_para.style.casefold()

    # 1. Strictest: Same Heading Style AND current_num + 1
    if is_heading and current_num is not None:
        for i in range(end_idx + 1, len(paragraphs)):
            real_p = document.paragraphs[i]
            if not _is_main_body(real_p): continue
            if paragraphs[i].style == end_para.style:
                p_text = _normalize_anchor(paragraphs[i].text)
                if chapter_number_from_text(p_text) == current_num + 1:
                    return i

    # 2. Same Heading Style (if it's a heading but maybe doesn't have a parseable number like "HASIL")
    if is_heading:
        for i in range(end_idx + 1, len(paragraphs)):
            real_p = document.paragraphs[i]
            if not _is_main_body(real_p): continue
            if paragraphs[i].style == end_para.style:
                return i

    # 3. Explicit "BAB " + Number
    if current_num is not None:
        for i in range(end_idx + 1, len(paragraphs)):
            real_p = document.paragraphs[i]
            if not _is_main_body(real_p): continue
            p_text = _normalize_anchor(paragraphs[i].text)
            if p_text.startswith("bab ") and chapter_number_from_text(p_text) == current_num + 1 and len(p_text.split()) < 15:
                return i

    # 4. Fallback: check if the end_anchor was an MS Word auto-numbered list item. If so, find the next item with the EXACT same list level.
    real_para = document.paragraphs[end_idx]
    if real_para._p.pPr is not None and real_para._p.pPr.numPr is not None:
        ilvl_val = real_para._p.pPr.numPr.ilvl.val if real_para._p.pPr.numPr.ilvl is not None else 0
        for i in range(end_idx + 1, len(paragraphs)):
            real_p = document.paragraphs[i]
            if not _is_main_body(real_p): continue
            i_pPr = real_p._p.pPr
            if i_pPr is not None and i_pPr.numPr is not None:
                i_ilvl = i_pPr.numPr.ilvl.val if i_pPr.numPr.ilvl is not None else 0
                if i_ilvl == ilvl_val:
                    return i

    # 5. Standard thesis major boundaries
    standard_headings = {
        "daftar pustaka", "daftar isi", "daftar gambar", "daftar tabel", 
        "daftar lampiran", "kata pengantar", "lembar pengesahan", "lampiran", "abstrak", "abstract",
        "riwayat hidup", "daftar singkatan"
    }
    for i in range(end_idx + 1, len(paragraphs)):
        real_p = document.paragraphs[i]
        if not _is_main_body(real_p): continue
        p_text = _normalize_anchor(paragraphs[i].text)
        if p_text in standard_headings or (len(p_text.split()) < 10 and any(p_text.startswith(h) for h in standard_headings)):
            return i
            
    return None


def add_section_breaks_for_command(document, command: DocumentCommand) -> None:
    """Insert breaks before range starts (if conflicting) and after range ends (to isolate)."""
    conflicts, _ = analyze_command(document, command)
    paragraphs = list_paragraphs(document, include_empty=True)
    
    # 1. Split ranges that currently share a section
    if conflicts:
        anchors_to_split = set()
        first_range_for_section: set[int] = set()
        for page_range in command.ranges:
            start_section = _find_anchor_section(paragraphs, page_range.start_anchor)
            if any(conflict.section_index == start_section for conflict in conflicts):
                if start_section in first_range_for_section:
                    anchors_to_split.add(page_range.start_anchor)
                else:
                    first_range_for_section.add(start_section)

        for anchor in anchors_to_split:
            paragraph_index = _find_anchor_paragraph_index(paragraphs, anchor)
            _insert_section_break_before(document, paragraphs, paragraph_index, anchor)

    # 2. Isolate the end of ranges.
    for page_range in command.ranges:
        if page_range.end_anchor == "__DOCUMENT_END__":
            continue
        try:
            end_idx = _find_anchor_paragraph_index(paragraphs, page_range.end_anchor)
            next_idx = _find_next_isolation_index(document, paragraphs, end_idx)
            
            if next_idx is not None:
                _insert_section_break_before(document, paragraphs, next_idx, f"auto-isolation after {page_range.end_anchor}")
        except ValueError:
            pass  # Anchor not found or ambiguous, skip isolation


def _insert_section_break_before(document, paragraphs, paragraph_index: int, anchor_name: str) -> None:
    if paragraph_index == 0:
        return  # Cannot insert before the first paragraph
    paragraph = document.paragraphs[paragraph_index]
    if paragraph._p.getparent().tag != qn("w:body"):
        return  # Cannot insert inside a table
    previous = document.paragraphs[paragraph_index - 1]
    paragraph_properties = previous._p.get_or_add_pPr()
    if paragraph_properties.sectPr is not None:
        return  # Already has a section break
    
    # We must find the current section's sectPr to copy it, maintaining margins etc.
    # We can do this by walking backward to find the previous sectPr, or just use the last section.
    # The safest way is to find the section index of the current paragraph.
    current_section_idx = paragraphs[paragraph_index].section_index
    current_section = document.sections[current_section_idx]
    new_sectPr = deepcopy(current_section._sectPr)
    
    # Strip hard-coded header/footer references so python-docx can properly unlink them later
    for ref in new_sectPr.findall(qn("w:headerReference")):
        new_sectPr.remove(ref)
    for ref in new_sectPr.findall(qn("w:footerReference")):
        new_sectPr.remove(ref)
        
    paragraph_properties.append(new_sectPr)


def _get_anchor_candidates(paragraphs, anchor: str):
    """Return paragraphs that best match *anchor*.

    Priority (highest wins):
    1. Anchor is "Bab N" or "Bab I" → match by chapter number, prefer headings.
    2. Anchor matches a heading whose text *starts with* the anchor (e.g. anchor="pendahuluan"
       matches heading "BAB I PENDAHULUAN").  Heading-only; must be unique per section.
    3. Exact full-text match (heading preferred).
    4. Partial substring match in headings (heading preferred).
    5. Fuzzy SequenceMatcher fallback (≥ 0.55).
    """
    normalized_anchor, aliases = _normalize_anchor_variants(anchor)

    # Reject junk anchors from misbehaving LLMs (e.g. "paragraph_index=1").
    if re.fullmatch(r"paragraph_index\s*=\s*\d+", normalized_anchor):
        return []

    # ── 1. Bab-number anchor ("bab 1", "bab i", "bab iii pendahuluan") ──
    chapter_number_match = re.match(r"bab\s+([0-9]+|[ivxlcdm]+)", normalized_anchor)
    if chapter_number_match:
        requested_number = parse_chapter_token(chapter_number_match.group(1))
        if requested_number is not None:
            chapter_matches = [
                paragraph
                for paragraph in paragraphs
                if chapter_number_from_text(_normalize_anchor(paragraph.text)) == requested_number
            ]
            heading_chapter_matches = [
                paragraph
                for paragraph in chapter_matches
                if "heading" in paragraph.style.casefold()
            ]
            if heading_chapter_matches:
                return heading_chapter_matches
            if chapter_matches:
                return chapter_matches

    variants = (normalized_anchor, *aliases)

    # ── 2. Heading starts-with match (anchor is a keyword inside a heading) ──
    heading_startswith = [
        paragraph
        for paragraph in paragraphs
        if "heading" in paragraph.style.casefold()
        and any(
            _normalize_anchor(paragraph.text).startswith(variant)
            or _normalize_anchor(paragraph.text).endswith(variant)
            or variant in _normalize_anchor(paragraph.text).split()
            for variant in variants
        )
    ]
    # Deduplicate: if all matches resolve to the same section, return them
    if heading_startswith:
        sections_hit = {p.section_index for p in heading_startswith}
        if len(sections_hit) == 1:
            return heading_startswith

    # ── 3. Exact full-text match ──
    # Prefer the exact, un-aliased text first: when a document contains both
    # "ABSTRAK" and "ABSTRACT", anchor "abstract" must hit the English heading,
    # not become ambiguous by also matching the Indonesian one.
    primary_exact = [
        paragraph
        for paragraph in paragraphs
        if _normalize_anchor(paragraph.text) == normalized_anchor
    ]
    heading_primary_exact = [p for p in primary_exact if "heading" in p.style.casefold()]
    if heading_primary_exact:
        return heading_primary_exact
    if primary_exact:
        return primary_exact
    exact_matches = [
        paragraph
        for paragraph in paragraphs
        if any(_normalize_anchor(paragraph.text) == variant for variant in variants)
    ]
    heading_exact = [p for p in exact_matches if "heading" in p.style.casefold()]
    if heading_exact:
        return heading_exact
    if exact_matches:
        return exact_matches

    # ── 4. Partial substring in headings ──
    partial_heading = [
        paragraph
        for paragraph in paragraphs
        if "heading" in paragraph.style.casefold()
        and any(variant in _normalize_anchor(paragraph.text) for variant in variants)
    ]
    if partial_heading:
        sections_hit = {p.section_index for p in partial_heading}
        if len(sections_hit) == 1:
            return partial_heading

    # Partial substring in all paragraphs
    partial_all = [
        paragraph
        for paragraph in paragraphs
        if any(variant in _normalize_anchor(paragraph.text) for variant in variants)
    ]
    if partial_all:
        return partial_all

    # ── 5. Fuzzy fallback ──
    scored = sorted(
        [
            (
                max(
                    SequenceMatcher(None, variant, _normalize_anchor(paragraph.text)).ratio()
                    for variant in variants
                ),
                paragraph,
            )
            for paragraph in paragraphs
        ],
        key=lambda item: item[0],
    )
    if not scored or scored[-1][0] < 0.55:
        return []
    best_score = scored[-1][0]
    return [paragraph for score, paragraph in scored if best_score - score < 0.03]


def _find_anchor_paragraph_index(paragraphs: list[ParagraphInfo], anchor: str) -> int:
    """Return the paragraph index to insert a section break before.

    When multiple candidates exist (e.g. anchor keyword appears in body text too)
    but exactly one heading candidate is present, use that heading.
    """
    normalized_anchor = _normalize_anchor(anchor)
    if normalized_anchor in {"awal", "awal dokumen", "pertama"}:
        return 0

    candidates = _get_anchor_candidates(paragraphs, anchor)
    if not candidates:
        raise ValueError(f"Cannot choose a unique paragraph for section break: {anchor}")
    # If multiple, prefer heading candidates
    heading_candidates = [p for p in candidates if "heading" in p.style.casefold()]
    if len(heading_candidates) == 1:
        return heading_candidates[0].index
    if len(candidates) == 1:
        return candidates[0].index
    # Still ambiguous – pick the first heading (earliest in document)
    if heading_candidates:
        return heading_candidates[0].index
    raise ValueError(f"Cannot choose a unique paragraph for section break: {anchor}")


def open_document(source: bytes):
    return load_document(source)


def inspect_document(document):
    return list_paragraphs(document, include_empty=True)


def _extract_section_settings(section, dom_font_name, dom_font_size) -> dict:
    """Read existing page-number settings from a section (pgNumType + header/footer PAGE fields)."""
    sec_settings = {
        "format": "none",
        "position": "bottom",
        "alignment": "center",
        "continue_previous": False,
        "start_number": 1,
        "first_page": None,
        "font_name": dom_font_name,
        "font_size": dom_font_size,
    }

    # Read pgNumType
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        fmt = pgNumType.get(qn("w:fmt"))
        if fmt in {"decimal", "lowerRoman", "upperRoman"}:
            sec_settings["format"] = fmt
        start = pgNumType.get(qn("w:start"))
        if start is not None:
            sec_settings["start_number"] = int(start)
            sec_settings["continue_previous"] = False
        else:
            sec_settings["continue_previous"] = True
            sec_settings["start_number"] = None

    # Helper to check PAGE field in header/footer
    def check_hf(hf):
        if hf.is_linked_to_previous:
            return None
        for p in hf.paragraphs:
            xml = p._p.xml
            if 'w:instrText' in xml and 'PAGE' in xml:
                if p.alignment == WD_ALIGN_PARAGRAPH.LEFT: return "left"
                if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT: return "right"
                return "center"
        return None

    footer_align = check_hf(section.footer)
    header_align = check_hf(section.header)

    if footer_align:
        sec_settings["position"] = "bottom"
        sec_settings["alignment"] = footer_align
        if sec_settings["format"] == "none": sec_settings["format"] = "decimal"
    elif header_align:
        sec_settings["position"] = "top"
        sec_settings["alignment"] = header_align
        if sec_settings["format"] == "none": sec_settings["format"] = "decimal"

    if section.different_first_page_header_footer:
        fp_footer = check_hf(section.first_page_footer)
        fp_header = check_hf(section.first_page_header)
        if fp_footer:
            sec_settings["first_page"] = {"position": "bottom", "alignment": fp_footer}
        elif fp_header:
            sec_settings["first_page"] = {"position": "top", "alignment": fp_header}

    return sec_settings


def apply_command(document, command: DocumentCommand) -> None:
    """Resolve semantic anchors, then apply page-number settings by section."""
    dom_font_name, dom_font_size = detect_dominant_font(document)
    
    if command.action == "clear_all_page_numbers":
        add_page_numbers(
            document,
            formats=["none"] * len(document.sections),
            positions=["bottom"] * len(document.sections),
            alignments=["center"] * len(document.sections),
            continue_previous=[False] * len(document.sections),
            first_page_settings=[None] * len(document.sections),
            start_numbers=[1] * len(document.sections),
        )
        return

    if command.action == "clear_page_number_ranges":
        paragraphs = list_paragraphs(document, include_empty=True)
        section_indices = set()
        for page_range in command.ranges:
            start_section = _find_anchor_section(paragraphs, page_range.start_anchor)
            end_section = _find_anchor_section(paragraphs, page_range.end_anchor)
            if start_section > end_section:
                raise ValueError("The clear range is reversed")
            section_indices.update(range(start_section, end_section + 1))
        clear_page_numbers_for_sections(document, section_indices)
        return

    add_section_breaks_for_command(document, command)
    paragraphs = list_paragraphs(document, include_empty=True)

    settings = [
        _extract_section_settings(sec, dom_font_name, dom_font_size)
        for sec in document.sections
    ]
    for page_range in command.ranges:
        range_settings = page_range.settings.model_dump()
        try:
            start_section = _find_anchor_section(paragraphs, page_range.start_anchor)
            
            if page_range.end_anchor == "__DOCUMENT_END__":
                end_section = len(document.sections) - 1
            else:
                end_idx = _find_anchor_paragraph_index(paragraphs, page_range.end_anchor)
                next_idx = _find_next_isolation_index(document, paragraphs, end_idx)
                
                if next_idx is not None:
                    end_section = paragraphs[next_idx - 1].section_index
                else:
                    end_section = len(document.sections) - 1
        except ValueError:
            continue

        for section_index in range(start_section, end_section + 1):
            sec_settings = dict(range_settings)
            # If a chapter spans multiple sections, only its first section gets the first_page override
            if section_index != start_section:
                sec_settings["first_page"] = None
            
            # Apply logic for continue_previous and start_number
            if section_index > start_section:
                sec_settings["continue_previous"] = True
                sec_settings["start_number"] = None
            else:
                sec_settings["continue_previous"] = page_range.settings.continue_previous
                sec_settings["start_number"] = (
                    page_range.settings.start_number
                    if not page_range.settings.continue_previous
                    else None
                )
                
            sec_settings["font_name"] = page_range.settings.font_name or dom_font_name
            sec_settings["font_size"] = page_range.settings.font_size or dom_font_size
            settings[section_index] = sec_settings

    add_page_numbers(
        document,
        formats=[section["format"] for section in settings],
        positions=[section["position"] for section in settings],
        alignments=[section["alignment"] for section in settings],
        continue_previous=[section["continue_previous"] for section in settings],
        first_page_settings=[section["first_page"] for section in settings],
        start_numbers=[section["start_number"] for section in settings],
        font_names=[section.get("font_name") for section in settings],
        font_sizes=[section.get("font_size") for section in settings],
    )


def _find_anchor_section(paragraphs, anchor: str) -> int:
    candidates = _get_anchor_candidates(paragraphs, anchor)
    if not candidates:
        raise ValueError(f"Could not find section anchor: {anchor}")
    matches = {paragraph.section_index for paragraph in candidates}
    if len(matches) > 1:
        raise ValueError(f"Section anchor is ambiguous: {anchor}")
    return next(iter(matches))


def _normalize_anchor(value: str) -> str:
    """Treat spaces, underscores, and hyphens consistently in AI anchors."""
    return " ".join(value.casefold().replace("_", " ").replace("-", " ").split())


# Cross-language anchor aliases so English spellings of common thesis sections
# resolve to their Indonesian heading in the document (and vice versa).
_ANCHOR_ALIASES = {
    "abstract": ("abstrak",),
    "abstrak": ("abstract",),
}


def _normalize_anchor_variants(value: str) -> tuple[str, tuple[str, ...]]:
    """Return (normalized_anchor, aliases) for fuzzy/partial matching.

    Useful when users reference "abstract" while the document heading says
    "ABSTRAK" (or the reverse), which previously fell through to a fuzzy match
    against body text and could land on the wrong chapter.
    """
    normalized = _normalize_anchor(value)
    return normalized, _ANCHOR_ALIASES.get(normalized, ())


def export_document(document) -> bytes:
    return to_bytes(document)

