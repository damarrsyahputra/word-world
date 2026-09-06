from __future__ import annotations

from io import BytesIO
from collections import Counter

from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

SUPPORTED_PAGE_NUMBER_FORMATS = {
    "none": "No page number",
    "decimal": "Decimal (1, 2, 3)",
    "lowerRoman": "Lower Roman (i, ii, iii)",
    "upperRoman": "Upper Roman (I, II, III)",
}
SUPPORTED_PAGE_NUMBER_POSITIONS = {
    "top": "Top",
    "bottom": "Bottom",
}


def add_page_numbers(
    document: DocumentType,
    alignment: str | None = None,
    formats: list[str] | None = None,
    positions: list[str] | None = None,
    alignments: list[str] | None = None,
    restart_at_one: bool = True,
    continue_previous: list[bool] | None = None,
    first_page_settings: list[dict[str, str] | None] | None = None,
    start_numbers: list[int | None] | None = None,
    font_names: list[str | None] | None = None,
    font_sizes: list[float | None] | None = None,
) -> None:
    """Configure dynamic PAGE fields independently for each document section."""
    if alignments is None:
        alignments = [alignment or "center"] * len(document.sections)
    if len(alignments) != len(document.sections):
        raise ValueError("An alignment is required for each section")
    if any(section_alignment not in {"left", "center", "right"} for section_alignment in alignments):
        raise ValueError("Alignment must be left, center, or right")
    if formats is None:
        formats = ["decimal"] * len(document.sections)
    if len(formats) != len(document.sections):
        raise ValueError("A page number format is required for each section")
    if positions is None:
        positions = ["bottom"] * len(document.sections)
    if len(positions) != len(document.sections):
        raise ValueError("A page number position is required for each section")
    if continue_previous is None:
        continue_previous = [False] * len(document.sections)
    if len(continue_previous) != len(document.sections):
        raise ValueError("A continuation setting is required for each section")
    if first_page_settings is None:
        first_page_settings = [None] * len(document.sections)
    if len(first_page_settings) != len(document.sections):
        raise ValueError("A first-page setting is required for each section")
    if start_numbers is None:
        start_numbers = [1] * len(document.sections)
    if len(start_numbers) != len(document.sections):
        raise ValueError("A start number is required for each section")

    if font_names is None: font_names = [None] * len(document.sections)
    if font_sizes is None: font_sizes = [None] * len(document.sections)
    for section, page_format, position, section_alignment, continues, first_page, start_number, f_name, f_size in zip(
        document.sections, formats, positions, alignments, continue_previous, first_page_settings, start_numbers, font_names, font_sizes
    ):
        if page_format not in SUPPORTED_PAGE_NUMBER_FORMATS:
            raise ValueError(f"Unsupported page number format: {page_format}")
        if position not in SUPPORTED_PAGE_NUMBER_POSITIONS:
            raise ValueError(f"Unsupported page number position: {position}")

        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        _remove_page_fields(section.header)
        _remove_page_fields(section.footer)
        _remove_page_fields(section.first_page_header)
        _remove_page_fields(section.first_page_footer)

        section_properties = section._sectPr
        page_number_type = section_properties.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            section_properties.append(page_number_type)
        page_number_type.set(qn("w:fmt"), page_format)
        if not continues:
            page_number_type.set(qn("w:start"), str(start_number or 1))
        elif continues:
            start = page_number_type.get(qn("w:start"))
            if start is not None:
                page_number_type.attrib.pop(qn("w:start"))

        section.different_first_page_header_footer = first_page is not None
        if page_format == "none":
            continue

        _add_page_field(section, position, section_alignment, False, f_name, f_size)
        if first_page is not None:
            if first_page.get("show", True):
                _add_page_field(section, first_page["position"], first_page["alignment"], first_page=True, font_name=f_name, font_size=f_size)


def _add_page_field(section, position: str, alignment: str, first_page: bool = False, font_name: str | None = None, font_size: float | None = None) -> None:
    if position not in SUPPORTED_PAGE_NUMBER_POSITIONS:
        raise ValueError(f"Unsupported page number position: {position}")
    if alignment not in {"left", "center", "right"}:
        raise ValueError(f"Unsupported page number alignment: {alignment}")
    if first_page:
        container = section.first_page_header if position == "top" else section.first_page_footer
    else:
        container = section.header if position == "top" else section.footer
    paragraph = container.paragraphs[0]
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]

    run = paragraph.add_run()
    if font_name:
        run.font.name = font_name
        rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def _remove_page_fields(container) -> None:
    """Remove paragraphs containing PAGE fields in a header/footer."""
    paragraphs_to_clear = []
    for instruction in container._element.iter(qn("w:instrText")):
        if "PAGE" not in (instruction.text or "").upper():
            continue
        parent = instruction.getparent()
        while parent is not None and parent.tag != qn("w:p"):
            parent = parent.getparent()
        if parent is not None and parent not in paragraphs_to_clear:
            paragraphs_to_clear.append(parent)
            
    for p_element in paragraphs_to_clear:
        # Clear all runs in the paragraph to remove the field and any cached text
        pPr = p_element.find(qn("w:pPr"))
        p_element.clear()
        if pPr is not None:
            p_element.append(pPr)


def clear_page_numbers_for_sections(document: DocumentType, section_indices: set[int]) -> None:
    """Remove page fields only from selected sections."""
    for section_index in section_indices:
        section = document.sections[section_index]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        _remove_page_fields(section.header)
        _remove_page_fields(section.footer)
        _remove_page_fields(section.first_page_header)
        _remove_page_fields(section.first_page_footer)
        page_number_type = section._sectPr.find(qn("w:pgNumType"))
        section.different_first_page_header_footer = False


def to_bytes(document: DocumentType) -> bytes:
    """Serialize the edited document without reconstructing its contents."""
    output = BytesIO()
    document.save(output)
    return output.getvalue()

def detect_dominant_font(document: DocumentType) -> tuple[str, float]:
    font_names = Counter()
    font_sizes = Counter()
    
    for paragraph in document.paragraphs[:50]:
        p_font = None
        p_size = None
        style = paragraph.style
        while style:
            if style.font.name and not p_font:
                p_font = style.font.name
            if style.font.size and not p_size:
                p_size = style.font.size.pt
            style = style.base_style
            
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            name = run.font.name or p_font
            size = run.font.size.pt if run.font.size else p_size
            
            if name:
                font_names[name] += len(run.text)
            if size:
                font_sizes[size] += len(run.text)
                
    normal_style = document.styles['Normal'] if 'Normal' in document.styles else None
    default_name = normal_style.font.name if normal_style and normal_style.font.name else "Times New Roman"
    default_size = normal_style.font.size.pt if normal_style and normal_style.font.size else 12.0

    dominant_name = font_names.most_common(1)[0][0] if font_names else default_name
    dominant_size = font_sizes.most_common(1)[0][0] if font_sizes else default_size
    
    return dominant_name, dominant_size
