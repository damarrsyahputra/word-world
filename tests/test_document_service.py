from __future__ import annotations

import pytest

from docx import Document

from app.document.reader import ParagraphInfo, list_paragraphs
from app.document.editor import add_page_numbers
from app.services.document_service import (
    _extract_section_settings,
    _find_anchor_paragraph_index,
    _find_next_isolation_index,
    _get_anchor_candidates,
    _normalize_anchor,
)


def _para(index: int, text: str, style: str = "Normal", section_index: int = 0) -> ParagraphInfo:
    return ParagraphInfo(index=index, text=text, style=style, section_index=section_index)


# ── Anchor normalization ────────────────────────────────────────────


@pytest.mark.parametrize(
    ("value", "expected"),
    [
        ("BAB I PENDAHULUAN", "bab i pendahuluan"),
        ("daftar_isi", "daftar isi"),
        ("lembar-pengesahan", "lembar pengesahan"),
        ("  Judul   Tugas  Akhir ", "judul tugas akhir"),
        ("", ""),
    ],
)
def test_normalize_anchor(value: str, expected: str):
    assert _normalize_anchor(value) == expected


# ── Anchor candidates ───────────────────────────────────────────────


def _thesis_paragraphs():
    return [
        _para(0, "Lembar Pengesahan", "Normal"),
        _para(1, "BAB I PENDAHULUAN", "Heading 1"),
        _para(2, "Latar Belakang", "Heading 2"),
        _para(3, "BAB II TINJAUAN PUSTAKA", "Heading 1"),
        _para(4, "Daftar Pustaka", "Normal"),
    ]


def test_get_anchor_candidates_matches_chapter_number():
    paragraphs = _thesis_paragraphs()
    candidates = _get_anchor_candidates(paragraphs, "bab 1")
    assert [p.index for p in candidates] == [1]


def test_get_anchor_candidates_matches_roman_chapter_number():
    paragraphs = _thesis_paragraphs()
    candidates = _get_anchor_candidates(paragraphs, "BAB II")
    assert [p.index for p in candidates] == [3]


def test_get_anchor_candidates_heading_startswith():
    paragraphs = _thesis_paragraphs()
    candidates = _get_anchor_candidates(paragraphs, "pendahuluan")
    assert [p.index for p in candidates] == [1]


def test_get_anchor_candidates_abstract_matches_abstrak_heading():
    # User says "abstract" (English) but the document heading is "ABSTRAK".
    paragraphs = [
        _para(0, "Lembar Pengesahan", "Heading 1"),
        _para(1, "ABSTRAK", "Heading 1"),
        _para(2, "The abstract is a summary of the thesis.", "Normal"),
        _para(3, "BAB I PENDAHULUAN", "Heading 1"),
    ]
    candidates = _get_anchor_candidates(paragraphs, "abstract")
    assert [p.index for p in candidates] == [1]


def test_get_anchor_candidates_abstrak_matches_abstract_heading():
    # Document heading is the English "ABSTRACT", user typed "abstrak".
    paragraphs = [
        _para(0, "ABSTRACT", "Heading 1"),
        _para(1, "BAB I PENDAHULUAN", "Heading 1"),
    ]
    candidates = _get_anchor_candidates(paragraphs, "abstrak")
    assert [p.index for p in candidates] == [0]


def test_get_anchor_candidates_both_headings_disambiguates_exact():
    # A real thesis contains BOTH "ABSTRAK" and "ABSTRACT". The literal anchor
    # must win over the alias so it is never reported as ambiguous.
    paragraphs = [
        _para(0, "LEMBAR PENGESAHAN", "Heading 1", 0),
        _para(1, "ABSTRAK", "Heading 1", 1),
        _para(2, "ABSTRACT", "Heading 1", 2),
        _para(3, "BAB I PENDAHULUAN", "Heading 1", 3),
    ]
    assert [p.index for p in _get_anchor_candidates(paragraphs, "abstract")] == [2]
    assert [p.index for p in _get_anchor_candidates(paragraphs, "abstrak")] == [1]


def test_get_anchor_candidates_exact_match():
    paragraphs = _thesis_paragraphs()
    candidates = _get_anchor_candidates(paragraphs, "daftar pustaka")
    assert [p.index for p in candidates] == [4]


def test_get_anchor_candidates_rejects_junk():
    paragraphs = _thesis_paragraphs()
    assert _get_anchor_candidates(paragraphs, "paragraph_index=1") == []


def test_get_anchor_candidates_returns_empty_for_missing():
    paragraphs = _thesis_paragraphs()
    assert _get_anchor_candidates(paragraphs, "tidak ada teks ini") == []


def test_find_anchor_paragraph_index_prefers_heading():
    paragraphs = [
        _para(0, "BAB I PENDAHULUAN", "Heading 1"),
        _para(1, "pembahasan pendahuluan singkat", "Normal"),
    ]
    assert _find_anchor_paragraph_index(paragraphs, "pendahuluan") == 0


def test_find_anchor_paragraph_index_start_of_document():
    paragraphs = _thesis_paragraphs()
    assert _find_anchor_paragraph_index(paragraphs, "awal") == 0


# ── Isolation BOUNDARY after front-matter anchors ─────────────────────


def _build_one_giant_section_doc():
    """Front matter + chapters share a single empty-section document (the exact
    failure mode of a thesis without section breaks between front matter and body)."""
    document = Document()
    order = [
        ("Lembar Pengesahan", "Normal"),
        ("ABSTRAK", "Normal"),
        ("ABSTRACT", "Normal"),
        ("BAB I\nPENDAHULUAN", "Heading 1"),
        ("BAB II\nTINJAUAN PUSTAKA", "Heading 1"),
        ("DAFTAR PUSTAKA", "Heading 1"),
    ]
    for text, style in order:
        p = document.add_paragraph(text)
        p.style = document.styles[style]
    return document


def test_isolation_finds_bab_1_after_abstract():
    document = _build_one_giant_section_doc()
    paragraphs = list_paragraphs(document, include_empty=True)
    abstract_idx = _find_anchor_paragraph_index(paragraphs, "abstract")
    next_idx = _find_next_isolation_index(document, paragraphs, abstract_idx)
    assert paragraphs[next_idx].text == "BAB I\nPENDAHULUAN"


def test_isolation_falls_back_to_standard_heading_when_no_bab():
    document = Document()
    for text, style in [
        ("ABSTRACT", "Normal"),
        ("KATA PENGANTAR", "Heading 1"),
        ("DAFTAR PUSTAKA", "Heading 1"),
    ]:
        p = document.add_paragraph(text)
        p.style = document.styles[style]
    paragraphs = list_paragraphs(document, include_empty=True)
    abstract_idx = _find_anchor_paragraph_index(paragraphs, "abstract")
    next_idx = _find_next_isolation_index(document, paragraphs, abstract_idx)
    assert paragraphs[next_idx].text == "KATA PENGANTAR"


# ── Section-settings reader (hoisted from apply_command) ────────────


def test_extract_section_settings_defaults_on_fresh_section():
    document = Document()
    settings = _extract_section_settings(document.sections[0], "Times New Roman", 12.0)
    assert settings["format"] == "none"
    assert settings["position"] == "bottom"
    assert settings["alignment"] == "center"
    assert settings["continue_previous"] is False
    assert settings["start_number"] == 1
    assert settings["first_page"] is None
    assert settings["font_name"] == "Times New Roman"
    assert settings["font_size"] == 12.0


def test_extract_section_settings_reads_existing_page_fields():
    document = Document()
    add_page_numbers(
        document,
        formats=["decimal"],
        positions=["bottom"],
        alignments=["center"],
    )
    settings = _extract_section_settings(document.sections[0], "Arial", 11.0)
    assert settings["format"] == "decimal"
    assert settings["position"] == "bottom"
    assert settings["alignment"] == "center"